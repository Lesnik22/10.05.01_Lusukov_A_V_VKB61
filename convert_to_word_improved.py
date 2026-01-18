#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Улучшенная конвертация диплома из Markdown в Word документ"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re
import sys

def setup_document_styles(doc):
    """Настройка стилей документа"""
    # Заголовок 1
    try:
        style = doc.styles['Heading 1']
    except:
        style = doc.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
    
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(16)
    font.bold = True
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style.paragraph_format.space_after = Pt(12)
    
    # Заголовок 2
    try:
        style = doc.styles['Heading 2']
    except:
        style = doc.styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
    
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    font.bold = True
    style.paragraph_format.space_after = Pt(10)
    
    # Заголовок 3
    try:
        style = doc.styles['Heading 3']
    except:
        style = doc.styles.add_style('Heading 3', WD_STYLE_TYPE.PARAGRAPH)
    
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    font.bold = True
    style.paragraph_format.space_after = Pt(8)
    
    # Обычный текст
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

def add_formatted_text(paragraph, text):
    """Добавление текста с форматированием"""
    if not text or not text.strip():
        return
    
    # Очистка от Markdown форматирования для простоты
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)  # Убираем **жирный**
    text = re.sub(r'\*([^*]+)\*', r'\1', text)  # Убираем *курсив*
    text = re.sub(r'`([^`]+)`', r'\1', text)  # Убираем `код`
    
    paragraph.add_run(text)

def process_markdown_to_docx(md_file, docx_file):
    """Конвертация Markdown в Word документ"""
    doc = Document()
    setup_document_styles(doc)
    
    # Настройка полей страницы
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(2)
        section.bottom_margin = Inches(2)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1)
    
    # Чтение Markdown файла
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    
    i = 0
    in_code_block = False
    code_block_lines = []
    in_table = False
    table_rows = []
    
    while i < len(lines):
        line = lines[i].rstrip()
        original_line = line
        
        # Проверка на код блоки
        if line.startswith('```'):
            if in_code_block:
                # Конец блока кода
                if code_block_lines:
                    para = doc.add_paragraph()
                    run = para.add_run('\n'.join(code_block_lines))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(11)
                    code_block_lines = []
                in_code_block = False
            else:
                # Начало блока кода
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue
        
        # Разделители
        if line.strip() == '---':
            if i > 0:  # Не добавляем пустую строку в начале
                doc.add_paragraph()
            i += 1
            continue
        
        # Пропуск пустых строк
        if not line.strip():
            i += 1
            continue
        
        # Заголовки
        if line.startswith('# '):
            # H1
            heading_text = line[2:].strip()
            para = doc.add_heading(heading_text, level=1)
            # Если это главный заголовок - центрируем
            if i < 10:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            # H2
            heading_text = line[3:].strip()
            doc.add_heading(heading_text, level=2)
        elif line.startswith('### '):
            # H3
            heading_text = line[4:].strip()
            doc.add_heading(heading_text, level=3)
        elif line.startswith('#### '):
            # H4
            heading_text = line[5:].strip()
            doc.add_heading(heading_text, level=4)
        # Маркированные списки
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            bullet_text = line.strip()[2:].strip()
            # Очистка от форматирования
            bullet_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', bullet_text)
            bullet_text = re.sub(r'\*([^*]+)\*', r'\1', bullet_text)
            para = doc.add_paragraph(bullet_text, style='List Bullet')
            # Обработка вложенных пунктов
            indent = len(line) - len(line.lstrip())
            if indent > 2:
                para.paragraph_format.left_indent = Inches(0.5)
        # Нумерованные списки
        elif re.match(r'^\d+\.\s', line.strip()):
            num_text = re.sub(r'^\d+\.\s', '', line.strip())
            num_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', num_text)
            para = doc.add_paragraph(num_text, style='List Number')
        # Таблицы
        elif '|' in line and line.strip().startswith('|'):
            # Начало или продолжение таблицы
            if not in_table:
                in_table = True
                table_rows = []
            
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells and not all(c.replace('-', '').strip() == '' for c in cells):  # Пропускаем разделитель
                table_rows.append(cells)
        else:
            # Конец таблицы
            if in_table and table_rows:
                # Создание таблицы
                headers = table_rows[0] if table_rows else []
                if headers and len(table_rows) > 1:
                    table = doc.add_table(rows=1, cols=len(headers))
                    table.style = 'Light Grid Accent 1'
                    
                    # Заголовки
                    for col_idx, header in enumerate(headers):
                        cell = table.rows[0].cells[col_idx]
                        cell.text = header
                        if cell.paragraphs[0].runs:
                            cell.paragraphs[0].runs[0].font.bold = True
                    
                    # Данные
                    for row_idx in range(1, len(table_rows)):
                        row_data = table_rows[row_idx]
                        if len(row_data) == len(headers):
                            row = table.add_row()
                            for col_idx, data in enumerate(row_data):
                                row.cells[col_idx].text = data
                
                in_table = False
                table_rows = []
            
            # Обычный текст
            text = line.strip()
            # Очистка от форматирования
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            text = re.sub(r'\*([^*]+)\*', r'\1', text)
            text = re.sub(r'`([^`]+)`', r'\1', text)
            text = re.sub(r'✅', '[OK]', text)
            text = re.sub(r'❌', '[FAIL]', text)
            
            if text:
                para = doc.add_paragraph(text)
        
        i += 1
    
    # Обработка последней таблицы, если она была открыта
    if in_table and table_rows:
        headers = table_rows[0] if table_rows else []
        if headers and len(table_rows) > 1:
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Light Grid Accent 1'
            
            for col_idx, header in enumerate(headers):
                cell = table.rows[0].cells[col_idx]
                cell.text = header
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.bold = True
            
            for row_idx in range(1, len(table_rows)):
                row_data = table_rows[row_idx]
                if len(row_data) == len(headers):
                    row = table.add_row()
                    for col_idx, data in enumerate(row_data):
                        row.cells[col_idx].text = data
    
    # Сохранение документа
    doc.save(docx_file)

if __name__ == '__main__':
    input_file = 'Диплом_Люсюков_Леха_Александрович.md'
    output_file = 'Диплом_Люсюков_Леха_Александрович.docx'
    
    try:
        print(f"Конвертация {input_file} в {output_file}...")
        process_markdown_to_docx(input_file, output_file)
        print(f"\n[OK] Конвертация завершена успешно!")
        print(f"[OK] Файл сохранен: {output_file}")
    except Exception as e:
        print(f"\n[ERROR] Ошибка при конвертации: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
