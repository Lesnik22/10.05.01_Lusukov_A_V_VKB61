#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Финальная конвертация объединенного диплома в Word"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def setup_document_styles(doc):
    """Настройка стилей документа для диплома"""
    # Обычный текст
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.first_line_indent = Inches(0.5)  # Красная строка
    
    # Заголовок 1
    try:
        style = doc.styles['Heading 1']
    except:
        style = doc.styles.add_style('Heading 1', 1)
    
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(16)
    font.bold = True
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style.paragraph_format.space_before = Pt(12)
    style.paragraph_format.space_after = Pt(12)
    
    # Заголовок 2
    try:
        style = doc.styles['Heading 2']
    except:
        style = doc.styles.add_style('Heading 2', 2)
    
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    font.bold = True
    style.paragraph_format.space_before = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    
    # Заголовок 3
    try:
        style = doc.styles['Heading 3']
    except:
        style = doc.styles.add_style('Heading 3', 3)
    
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    font.bold = True
    style.paragraph_format.space_before = Pt(10)
    style.paragraph_format.space_after = Pt(6)

def process_markdown_to_docx(md_file, docx_file):
    """Конвертация Markdown в Word с правильной структурой для диплома"""
    doc = Document()
    setup_document_styles(doc)
    
    # Настройка полей страницы (стандарт для дипломов)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(2)
        section.bottom_margin = Inches(2)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1)
    
    # Чтение Markdown файла
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    
    i = 0
    in_code_block = False
    code_block_lines = []
    in_table = False
    table_rows = []
    
    while i < len(lines):
        line = lines[i].rstrip()
        original_line = line
        
        # Код блоки
        if line.startswith('```'):
            if in_code_block:
                # Конец блока кода
                if code_block_lines:
                    para = doc.add_paragraph()
                    run = para.add_run('\n'.join(code_block_lines))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(11)
                    code_block_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue
        
        # Разделители
        if line.strip() == '---':
            if i > 0:
                doc.add_page_break()
            i += 1
            continue
        
        # Пустые строки
        if not line.strip():
            i += 1
            continue
        
        # Заголовки
        if line.startswith('# '):
            heading_text = line[2:].strip()
            # Убираем форматирование из заголовков
            heading_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', heading_text)
            
            para = doc.add_heading(heading_text, level=1)
            # Титульный лист центрируем
            if i < 30:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            heading_text = line[3:].strip()
            heading_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', heading_text)
            doc.add_heading(heading_text, level=2)
        elif line.startswith('### '):
            heading_text = line[4:].strip()
            heading_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', heading_text)
            doc.add_heading(heading_text, level=3)
        elif line.startswith('#### '):
            heading_text = line[5:].strip()
            heading_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', heading_text)
            doc.add_heading(heading_text, level=4)
        # Маркированные списки
        elif line.strip().startswith('- ') or (line.strip().startswith('* ') and not line.strip().startswith('**')):
            bullet_text = line.strip()[2:].strip()
            # Очистка от форматирования
            bullet_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', bullet_text)
            bullet_text = re.sub(r'\*([^*]+)\*', r'\1', bullet_text)
            bullet_text = re.sub(r'✅', '[OK]', bullet_text)
            bullet_text = re.sub(r'❌', '[FAIL]', bullet_text)
            
            para = doc.add_paragraph(bullet_text, style='List Bullet')
            # Вложенные списки
            indent = len(line) - len(line.lstrip())
            if indent > 2:
                para.paragraph_format.left_indent = Inches(0.5)
        # Нумерованные списки
        elif re.match(r'^\d+\.\s', line.strip()):
            num_text = re.sub(r'^\d+\.\s', '', line.strip())
            num_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', num_text)
            num_text = re.sub(r'\*([^*]+)\*', r'\1', num_text)
            para = doc.add_paragraph(num_text, style='List Number')
        # Таблицы
        elif '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            # Пропускаем разделитель таблицы
            if cells and not all(c.replace('-', '').replace(':', '').strip() == '' for c in cells):
                table_rows.append(cells)
        else:
            # Конец таблицы
            if in_table and table_rows and len(table_rows) > 1:
                headers = table_rows[0]
                if headers:
                    table = doc.add_table(rows=1, cols=len(headers))
                    table.style = 'Light Grid Accent 1'
                    
                    # Заголовки
                    for col_idx, header in enumerate(headers):
                        cell = table.rows[0].cells[col_idx]
                        cell.text = header
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.bold = True
                    
                    # Данные
                    for row_idx in range(1, len(table_rows)):
                        row_data = table_rows[row_idx]
                        if len(row_data) == len(headers):
                            row = table.add_row()
                            for col_idx, data in enumerate(row_data):
                                row.cells[col_idx].text = data
                
                in_table = False
                table_rows = []
            
            # Обычный текст
            text = line.strip()
            
            # Очистка от форматирования, но сохраняем структуру
            # Заменяем жирный текст просто на текст
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            # Убираем курсив (или оставляем, если нужно)
            text = re.sub(r'\*([^*]+)\*', r'\1', text)
            # Убираем код
            text = re.sub(r'`([^`]+)`', r'\1', text)
            # Заменяем эмодзи
            text = re.sub(r'✅', '[OK]', text)
            text = re.sub(r'❌', '[FAIL]', text)
            
            if text and not text.startswith('|'):
                para = doc.add_paragraph(text)
        
        i += 1
    
    # Обработка последней таблицы
    if in_table and table_rows and len(table_rows) > 1:
        headers = table_rows[0]
        if headers:
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Light Grid Accent 1'
            
            for col_idx, header in enumerate(headers):
                cell = table.rows[0].cells[col_idx]
                cell.text = header
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.bold = True
            
            for row_idx in range(1, len(table_rows)):
                row_data = table_rows[row_idx]
                if len(row_data) == len(headers):
                    row = table.add_row()
                    for col_idx, data in enumerate(row_data):
                        row.cells[col_idx].text = data
    
    # Сохранение документа
    doc.save(docx_file)

if __name__ == '__main__':
    input_file = 'Диплом_Люсюков_Леха_Александрович.md'
    output_file = 'Диплом_Люсюков_Леха_Александрович_ИТОГОВЫЙ.docx'
    
    try:
        print(f"Конвертация {input_file} в {output_file}...")
        process_markdown_to_docx(input_file, output_file)
        print(f"\n[OK] Конвертация завершена успешно!")
        print(f"[OK] Файл сохранен: {output_file}")
        print("\nПримечание: Откройте файл в Word и:")
        print("1. Добавьте автосодержание (Ссылки -> Оглавление)")
        print("2. Проверьте нумерацию страниц")
        print("3. Отформатируйте титульный лист согласно требованиям")
        print("4. Добавьте информацию из отчета о практике в соответствующие разделы")
    except Exception as e:
        print(f"\n[ERROR] Ошибка при конвертации: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
